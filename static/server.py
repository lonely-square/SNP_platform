from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx import Document
from docx.shared import Inches
import socket

from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne

# refer to docx.oxml.shape.CT_Inline
class CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')
 
    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor
 
    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor
    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'                    
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapNone/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % ( nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y) )
        )
# refer to docx.parts.story.BaseStoryPart.new_pic_inline
def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    """Return a newly-created `w:anchor` element.
    The element contains the image specified by *image_descriptor* and is scaled
    based on the values of *width* and *height*.
    """
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename    
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)
# refer to docx.text.run.add_picture
def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page.
    """
    run = p.add_run()
    anchor = new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
    run._r.add_drawing(anchor)
# refer to docx.oxml.__init__.py
register_element_cls('wp:anchor', CT_Anchor)



def createSickTable(sheet,sickNum):
    nlen=sheet.nrows
    clen=sheet.ncols

    table = document.add_table(rows=1,cols=2,style='Colorful Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  #表格对齐
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '病名'
    hdr_cells[1].text = '患病几率'
    
    hdr_cells[0].font = Pt(24)
    hdr_cells[0].color = RGBColor(0,0,0)

    for i in range(nlen-1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(sheet2.cell(i+1,0).value)
        row_cells[1].text = str(round(sheet2.cell(i+1,sickNum).value,3))


def sickerReport(sheet2,sickNum):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 靠右
    document.add_heading('第{:}个患者的检测报告'.format(sickNum+1), 0).bold = True  # 文档标题

    """p = document.add_paragraph('A plain paragraph having some')  # 段落
    p.add_run('bold').bold = True  # 黑体
    p.add_run(' and some ')
    p.add_run('italic.').italic = True  # 斜体
    """
    document.add_heading('一.患病几率', level=1)  # 一级标题
    
    passage="现患率也称患病率，与发病率相同，同样用分数的大小来反映患病率的高低。与发病率不同的是，患病率的分子为特定时间一定人群中某病新旧病例数，不管它是新发病还是旧病，只要是特定时间内疾病尚未痊愈，就记为病例数，然而发病率的分子为一定期间暴露人群中新病例人数，暴露人群中任何人新发生某疾病都称为新病例。"
    
    p1=document.add_paragraph()

    text1 = p1.add_run(passage)
    
    text1.font.size = Pt(10)                # 字体大小
    text1.bold = True                    # 字体是否加粗
    text1.font.name = 'Times New Roman'           # 控制是西文时的字体
    text1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 控制是中文时的字体


    createSickTable(sheet2,sickNum+1)

    document.add_page_break()


def createReport():
    # 创建文档对象
    pythoncom.CoInitialize()
    document = Document()


    document.sections[0].page_width = Cm(21)
    document.sections[0].page_height = Cm(29.7)
    #第一页
    p = document.add_paragraph()
    add_float_picture(p, 'test.png', width=Cm(21), pos_x=Pt(0), pos_y=Pt(0))
    document.add_page_break()


    document.add_section()
    document.sections[1].page_width = Cm(21)
    document.sections[1].page_height = Cm(29.7)
    #修改页眉
    header = document.sections[1].header  # 获取第一个节的页眉
    header.is_linker_to_previous = False
    paragraph = header.paragraphs[0]  # 获取页眉的第一个段落
    paragraph.add_run('上海大学')  # 添加页面内容



    clen=sheet2.ncols

    for i in range(clen-1):
        sickerReport(sheet2,i)




    # 时间路径
    t = datetime.datetime.now().strftime('%Y%m%d%H%M')
    dname = t + "demo.docx"  # word文档文件名
    doc_path = os.path.join(doc_base_path, dname)  # word生成路径
    # 生成word文档
    document.save(doc_path)



    pname = t + "demo.pdf"  # pdf名
    pdf_path = os.path.join(pdf_base_path, pname)  # pdf路径


    # 创建空pdf
    c = canvas.Canvas(pdf_path)
    c.showPage()
    c.save()


    # 判断路径是否存在,存在的话，将.docx文档转为.pdf文档
    if os.path.exists(doc_base_path) and os.path.exists(pdf_base_path):
        if os.path.isfile(doc_path) and os.path.isfile(pdf_path):  # 判断是否存在该word和pdf文件
            from docx2pdf import convert

            convert(doc_path, pdf_path)  # word转pdf
        else:
            print("文档不存在")
    else:
        print("路径不存在")



    
import xlrd
# 打开文件
workbook = xlrd.open_workbook(r'Score_41.xlsx')
# 获取所有sheet
print(workbook.sheet_names())
#获取sheet2
sheet2_name= workbook.sheet_names()[0]
print(sheet2_name)
# 根据sheet索引或者名称获取sheet内容
sheet2 = workbook.sheet_by_name('Sheet1')

#socket应用TCP（由服务器和和客户端之分） UDP（不回传丢失数据包）
while True:
    with socket.socket(socket.AF_INET,socket.SOCK_STREAM) as s:  #s是ServerSocket，c是clientSocker  with离开with块以后自动销毁对象
        #AF_INET标志用的是IPv4,SOCK_STREAM申明TCP协议
        s.bind(("0.0.0.0",1234))    #关联网卡，0.0.0.0表示任意网卡可以连接，1234表示端口
        s.listen()                  #设施监听
        c,addr =s.accept()          #返回新socket c和链接客户端IP地址
        with c:
            print(addr,"connected.")      
            while True:
                data = c.recv(1024)  #接收信息，最大长度1024字节
                if not data:
                    break
                c.sendall(createReport())


 
